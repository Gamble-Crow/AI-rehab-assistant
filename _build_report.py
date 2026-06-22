# -*- coding: utf-8 -*-
"""Bao cao TTCS v2 - cap nhat bam ban code moi nhat (UI moi, camera day so lieu, deploy offline)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\Admin\Desktop\BAO_CAO_TTCS_v2_moi.docx"
doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Arial"; normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
for i, sz in [(1, 18), (2, 15), (3, 13)]:
    st = doc.styles["Heading %d" % i]
    st.font.name = "Arial"; st.font.size = Pt(sz); st.font.bold = True
    st.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)

HDR_BG = "D5E1F0"; CODE_BG = "F2F2F2"

def H1(t): doc.add_heading(t, level=1)
def H2(t): doc.add_heading(t, level=2)
def H3(t): doc.add_heading(t, level=3)

def P(t="", bold=False, italic=False, align=None, size=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if align == "c": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def B(t):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(t); return p
def N(t):
    p = doc.add_paragraph(style="List Number"); p.add_run(t); return p

def PD(name, desc):
    p = doc.add_paragraph(); r = p.add_run(name); r.bold = True
    p.add_run(" — " + desc); return p

def LBL(label, text):
    p = doc.add_paragraph(); r = p.add_run(label + " "); r.bold = True
    p.add_run(text); return p

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor); tcPr.append(shd)

def _borders(tbl):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge); e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "BBBBBB"); borders.append(e)
    tbl._tbl.tblPr.append(borders)

def CODE(code, caption=None):
    if caption:
        cp = doc.add_paragraph(); rr = cp.add_run(caption); rr.bold = True; rr.font.size = Pt(10.5)
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]; _shade(cell, CODE_BG); _borders(tbl)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in code.strip("\n").split("\n"):
        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas"); r.font.size = Pt(9)
    doc.add_paragraph()

def TABLE(headers, rows, widths=None, caption=None):
    if caption:
        cp = doc.add_paragraph(); rr = cp.add_run(caption); rr.bold = True; rr.font.size = Pt(10.5)
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; _borders(tbl)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], HDR_BG); run = hdr[i].paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10.5)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(10.5)
    if widths:
        for i, w in enumerate(widths):
            for row in tbl.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph()

def PB(): doc.add_page_break()

# =====================================================================
# TRANG BIA
# =====================================================================
P("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", bold=True, align="c", size=14)
P("KHOA CÔNG NGHỆ THÔNG TIN", bold=True, align="c", size=14)
P(""); P("")
P("BÁO CÁO THỰC TẬP CƠ SỞ", bold=True, align="c", size=20)
P("ĐỀ TÀI: Xây dựng AI hỗ trợ trong tập phục hồi chức năng cơ bản cho bệnh nhân chấn thương thể thao",
  bold=True, align="c", size=15)
P(""); P("")
P("Giảng viên hướng dẫn : Đặng Hoàng Long", align="c")
P("Nhóm : 03", align="c"); P("")
TABLE(["Tên thành viên", "MSSV"],
      [["Bùi Trí Hiếu", "B23DCVT150"], ["Nguyễn Đức Dương", "B23DCVT114"],
       ["Đinh Khánh Sơn", "B23DCDT223"], ["Phạm Dũng", "B23DCCE024"]],
      widths=[3.8, 2.5])
P("Hà Nội - 2026", bold=True, align="c")
PB()
P("BẢNG PHÂN CÔNG CÔNG VIỆC", bold=True, align="c", size=14); P("")
TABLE(["Thành viên", "20/3", "3/5", "23/6"],
      [["Bùi Trí Hiếu", "camera (ver 1)", "brain engine; camera (ver 2)", "soát lại, sửa bug"],
       ["Nguyễn Đức Dương", "camera (ver 1); báo cáo", "báo cáo; database (video ver 1); kết nối", "làm slide; báo cáo"],
       ["Đinh Khánh Sơn", "database (ver 1)", "database (ver 2); audio (ver 2)", "database (video ver 2); soát lại, sửa bug"],
       ["Phạm Dũng", "audio (ver 1)", "audio (ver 2); UI", "làm slide; báo cáo"]],
      widths=[1.2, 1.5, 2.1, 1.6])
PB()

# =====================================================================
# MUC LUC
# =====================================================================
H1("Mục lục")
toc = [
    "Lời mở đầu", "I. Giới thiệu đề tài", "    1. Đối tượng",
    "    2. Sự khác biệt giữa tập gym thường và tập phục hồi chức năng",
    "    3. Tham khảo", "    4. Điểm nổi bật của dự án",
    "II. Cấu trúc tổng thể", "    1. Sơ đồ hoạt động",
    "    2. Các mảng kiến thức cốt lõi (Là gì? / Tại sao dùng? / Áp dụng?)",
    "    3. Sơ đồ chi tiết các module (flowchart)",
    "    4. Database (sơ đồ thực thể - quan hệ)",
    "III. Thiết kế và cài đặt",
    "    1. Database", "    2. Camera", "    3. Audio processing",
    "    4. Brain engine", "    5. API & UI bridge",
    "IV. Sản phẩm", "    1. Triển khai (deploy) & yêu cầu cấu hình",
    "    2. Thông báo / phản hồi cho người dùng", "    3. Dữ liệu & cách tạo bài tập",
    "    4. Phương pháp đánh giá", "    5. Demo sản phẩm", "V. Kết luận",
]
for t in toc:
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(12); p.paragraph_format.space_after = Pt(2)
PB()

# =====================================================================
# LOI MO DAU
# =====================================================================
H1("Lời mở đầu")
P("Trong thời đại công nghệ số phát triển mạnh mẽ, trí tuệ nhân tạo (AI) đang ngày càng được "
  "ứng dụng rộng rãi trong nhiều lĩnh vực của đời sống, đặc biệt là trong y tế. Một trong "
  "những hướng ứng dụng tiềm năng của AI là hỗ trợ quá trình phục hồi chức năng cho bệnh nhân "
  "sau chấn thương, phẫu thuật hoặc các bệnh lý ảnh hưởng đến khả năng vận động. Việc nghiên "
  "cứu và xây dựng các hệ thống AI hỗ trợ phục hồi chức năng cơ bản không chỉ giúp nâng cao "
  "hiệu quả điều trị mà còn góp phần giảm tải cho đội ngũ y tế và tạo điều kiện cho bệnh nhân "
  "luyện tập thuận tiện hơn.")
P("Xuất phát từ ý nghĩa thực tiễn đó, nhóm chúng em thực hiện đề tài “Xây dựng hệ thống "
  "AI hỗ trợ trong phục hồi chức năng cơ bản cho bệnh nhân chấn thương thể thao”, tập "
  "trung vào hai nhóm bệnh điển hình là đứt dây chằng chéo trước (ACL) ở khớp gối và hội "
  "chứng “Tennis Elbow” ở khớp khuỷu tay, nhằm tìm hiểu các phương pháp ứng dụng trí "
  "tuệ nhân tạo trong việc theo dõi, hướng dẫn và đánh giá quá trình luyện tập phục hồi của "
  "bệnh nhân.")
P("Trong quá trình thực hiện, chúng em đã nhận được sự hướng dẫn và hỗ trợ tận tình từ giảng "
  "viên Đặng Hoàng Long và chị trợ giảng. Chúng em xin gửi lời cảm ơn chân thành đến thầy và "
  "chị vì đã cung cấp những kiến thức quý báu, định hướng nghiên cứu và giúp chúng em hoàn "
  "thành báo cáo này.")
P("Mặc dù đã cố gắng trong quá trình thực hiện, tuy nhiên do kiến thức và kinh nghiệm còn hạn "
  "chế nên báo cáo khó tránh khỏi những thiếu sót. Nhóm chúng em rất mong nhận được những ý "
  "kiến đóng góp quý báu từ thầy để bài báo cáo được hoàn thiện hơn.")
PB()

# =====================================================================
# I. GIOI THIEU DE TAI
# =====================================================================
H1("I. Giới thiệu đề tài")
H2("1. Đối tượng")
P("Đối tượng hướng đến là những bệnh nhân chấn thương do hoạt động thể thao và các ca sau "
  "phẫu thuật khớp nhưng đã hồi phục đủ để bắt đầu tự tập thêm ở nhà, ví dụ như đứt dây chằng "
  "đầu gối, tràn dịch khớp, cứng khuỷu tay sau bó bột...")
P("Mục tiêu của các bài tập được lựa chọn:")
B("Duy trì biên độ hoạt động của khớp (giai đoạn đầu).")
B("Phục hồi và tăng dần sức mạnh cũng như ổn định khớp (giai đoạn giữa).")
B("Đưa khả năng vận động của khớp về như trước chấn thương (giai đoạn cuối).")
P("Nhóm tập trung vào hai nhóm đối tượng cụ thể:")
B("Bệnh nhân rách dây chằng chéo trước đầu gối (ACL) sau 6 tuần (bắt đầu tập phục hồi). Bộ "
  "phận tập là đầu gối.")
B("Bệnh nhân hội chứng “Tennis Elbow” (viêm cầu lồi ngoài xương cánh tay), khớp tập là "
  "khuỷu tay; thời điểm bắt đầu là sau 2-4 tuần khi bệnh nhân sinh hoạt hằng ngày mà không "
  "còn đau.")
H2("2. Sự khác biệt giữa tập gym thường và tập phục hồi chức năng")
P("Tập gym thông thường nhằm rèn luyện sức khỏe, tăng sức chịu đựng và khối cơ. Đau là điều "
  "bình thường khi tập gym do đây là tổn thương cơ dạng vi mô mà cơ thể có thể tự sửa chữa, "
  "đi kèm là cơ to và khỏe hơn. Người tập thường dùng tạ để tăng cường độ và có thể tự chọn "
  "lịch tập, thậm chí tập hằng ngày.")
P("Tập phục hồi chức năng nhằm khôi phục lại chức năng vận động sau chấn thương. Do đó khi "
  "bệnh nhân có dấu hiệu đau hay quá sức thì cần dừng tập hoặc giảm cường độ. Cường độ phải "
  "tăng chậm và kiểm soát chặt chẽ, đồng thời linh động theo khả năng phục hồi của từng người. "
  "Bên cạnh đó còn nhiều ràng buộc an toàn, ví dụ biên độ khớp không được vượt quá biên độ "
  "chỉ định: ACL mới mổ không được gập quá 90° trong 6 tuần đầu.")
P("Chính sự khác biệt này là lý do hệ thống cần đo góc khớp chính xác, phát hiện tín hiệu đau "
  "và điều chỉnh cường độ tự động - thay vì chỉ đếm số lần như app tập gym.")
H2("3. Tham khảo")
H3("a) Artificial Intelligence in the Management of Anterior Cruciate Ligament Injuries")
P("Bài báo nghiên cứu vai trò của AI trong tối ưu hóa kết quả cho bệnh nhân chấn thương ACL, "
  "chỉ ra 4 ứng dụng chính: chẩn đoán hình ảnh, dự đoán nguy cơ chấn thương, lập kế hoạch "
  "phẫu thuật, và phục hồi chức năng kèm đánh giá kết quả. Trong phục hồi chức năng, cần "
  "computer vision để theo dõi cử động của bệnh nhân tại nhà, kết hợp machine learning để dự "
  "đoán kết quả và mạng nơ-ron (ANN) để đánh giá nguy cơ tái chấn thương.")
P("Phân tích & áp dụng: Bài báo khẳng định tính cấp thiết của việc giám sát quá trình tập "
  "bằng thị giác máy tính và đo biên độ vận động (ROM). → Hệ thống của nhóm chọn hướng "
  "computer vision đo góc khớp làm lõi, và thiết kế “bộ não” đánh giá - điều chỉnh "
  "cường độ (thay cho mô hình dự đoán phức tạp vốn cần dữ liệu lâm sàng lớn mà nhóm không có).")
H3("b) Effectiveness of Digital Health Interventions for Rehabilitation of Elbow Injuries and "
   "Disorders: A Systematic Review")
P("Bài tổng quan đánh giá các công cụ kỹ thuật số trong cải thiện chức năng khớp khuỷu so với "
  "trị liệu thông thường, trên bệnh nhân sau phẫu thuật khuỷu, gãy xương vùng khuỷu hoặc "
  "“Tennis Elbow”. Ba kết quả chính: cải thiện ROM nhờ mô phỏng hình ảnh và nhận diện "
  "khớp qua camera; giảm đau và cải thiện chức năng; tương tác thời gian thực. Bài báo nhấn "
  "mạnh tính đối xứng tay thuận/nghịch và việc tập nhiều lần cường độ nhẹ hiệu quả hơn tập nặng.")
P("Phân tích & áp dụng: → Hệ thống cần đếm số lần (rep) và phản hồi tức thì để tạo động "
  "lực; cần hỗ trợ cả khuỷu tay; cần cơ chế tăng cường độ chậm - đây chính là cơ sở cho thuật "
  "toán “tăng +1, giảm mạnh” và chống dao động trong brain engine.")
H3("c) Curovate (ứng dụng thương mại)")
P("Curovate là một trong những ứng dụng phục hồi chức năng chuyên sâu, uy tín về ACL, thiết "
  "kế theo các giao thức lâm sàng chuẩn (chia lộ trình theo giai đoạn, đo ROM, cá nhân hóa "
  "cao). Ứng dụng dùng cloud để lưu dữ liệu, dùng thị giác máy tính đo góc khớp qua camera "
  "điện thoại.")
P("Phân tích & khoảng trống: Curovate còn nhiều hạn chế - ít tập trung các chấn thương khác, "
  "chưa theo dõi trực tiếp quá trình tập, chưa hỗ trợ chuyên sâu tại Việt Nam, và phải trả "
  "phí hằng tháng. → Đây là khoảng trống mà dự án nhắm tới: xử lý cục bộ (không cloud, "
  "bảo mật), miễn phí, theo dõi trực tiếp theo thời gian thực, và dễ mở rộng sang khớp khác.")
H3("d) Nguồn kỹ thuật của các mô hình AI sử dụng")
B("MediaPipe Pose / BlazePose (Google) - mô hình ước lượng tư thế, cho 33 điểm mốc cơ thể "
  "kèm tọa độ 3D, chạy thời gian thực trên CPU.")
B("YAMNet / AudioSet (Google) - mô hình phân loại 521 lớp âm thanh, dùng để nhận diện tiếng "
  "kêu/rên (không lời).")
B("Whisper / faster-whisper (OpenAI) - mô hình nhận dạng giọng nói (ASR), dùng để bắt các từ "
  "tiếng Việt thể hiện đau.")
P("Phân tích: cả ba đều là mô hình tiền huấn luyện (pre-trained), miễn phí, chạy được trên "
  "máy cá nhân cấu hình phổ thông - phù hợp ràng buộc “chi phí thấp, cục bộ” của đề tài.")
P("Tổng hợp → yêu cầu thiết kế rút ra từ khảo sát:", bold=True)
N("Đo góc khớp (ROM) chính xác bằng thị giác máy tính, không cần cảm biến chuyên dụng.")
N("Đếm rep và phản hồi tức thì để tạo động lực cho bệnh nhân.")
N("Phát hiện tín hiệu đau để đảm bảo an toàn.")
N("Tự điều chỉnh cường độ theo khả năng phục hồi, tăng chậm - giảm an toàn.")
N("Chạy cục bộ, chi phí thấp, bảo mật dữ liệu, dễ mở rộng sang nhiều khớp.")
H2("4. Điểm nổi bật của dự án")
P("Khác với ứng dụng thương mại như Curovate, dự án xử lý hoàn toàn thời gian thực ngay tại "
  "máy cục bộ (offline sau lần tải model đầu tiên), nhờ đó bảo mật dữ liệu tuyệt đối. Yêu cầu "
  "phần cứng tối thiểu - chỉ cần một thiết bị cá nhân có webcam và micro - giúp bệnh nhân dễ "
  "tiếp cận.")
B("Không chỉ tập trung một bộ phận mà dễ mở rộng sang các khớp khác; tiềm năng liên kết với "
  "bệnh viện, giảm chi phí và thời gian cho bệnh nhân, giảm tải cho cơ sở y tế.")
B("Tính mở và tùy biến: bác sĩ / kỹ thuật viên có thể chỉnh ngưỡng góc tập ngay trong "
  "database; hệ thống tự đưa ra gợi ý điều chỉnh cường độ theo khả năng hồi phục của từng "
  "bệnh nhân.")
B("Video hướng dẫn mẫu cho từng bài được lưu trong database và phát trực tiếp trên giao diện "
  "song song với khung camera theo dõi động tác.")
B("Dễ dàng chuyển từ bài đếm số lần (rep) sang bài đếm thời gian giữ nguyên tư thế.")
PB()

# =====================================================================
# II. CAU TRUC TONG THE
# =====================================================================
H1("II. Cấu trúc tổng thể")
H2("1. Sơ đồ hoạt động")
P("Quy trình một buổi tập: bệnh nhân đăng nhập (hoặc tạo hồ sơ mới) → chọn khớp và bài "
  "tập → đặt số rep mục tiêu → bắt đầu buổi tập. Khi buổi tập chạy, hai luồng xử lý "
  "song song: luồng camera (đo góc khớp, đếm rep) và luồng micro (phát hiện tiếng kêu đau). "
  "Kết quả được cập nhật tức thì lên giao diện. Khi kết thúc, dữ liệu buổi tập được đưa vào "
  "“bộ não” (brain engine) để chấm điểm và gợi ý tăng/giảm/giữ cường độ; bệnh nhân đồng "
  "ý hoặc từ chối, sau đó hệ thống lưu lại.")
CODE("""
@startuml
start
:Đăng nhập / tạo bệnh nhân;
:Chọn khớp + bài tập;
:Đặt số rep mục tiêu;
:Bắt đầu buổi tập;
fork
  :Luồng Camera (MediaPipe Pose 3D)
   - đo góc khớp -> FSM đếm rep;
fork again
  :Luồng Micro (YAMNet + Whisper)
   - đếm số lần kêu đau;
end fork
:Cập nhật rep / cảnh báo đau theo thời gian thực lên UI;
:Kết thúc buổi tập;
:Brain Engine chấm điểm (đau, tốc độ, ROM, hoàn thành);
:Đề xuất tăng / giảm / giữ + lý do;
:Bệnh nhân đồng ý / từ chối;
:Lưu session_log + cập nhật cường độ;
stop
@enduml
""", caption="Sơ đồ hoạt động tổng thể (PlantUML - activity diagram):")
H2("2. Các mảng kiến thức cốt lõi của hệ thống")
P("Mỗi mảng được trình bày theo ba câu hỏi: Là gì? – Tại sao chọn (so sánh các phương "
  "án)? – Áp dụng thế nào?", italic=True)
H3("2.1. Thị giác máy tính (Computer Vision) & Toán học vector")
P("Là gì? Ước lượng tư thế (Pose Estimation) là kỹ thuật xác định vị trí các khớp trên cơ "
  "thể từ ảnh/video. Kết hợp hình học vector, ta tính được góc của khớp cần theo dõi.")
P("Tại sao chọn MediaPipe Pose? So sánh các phương án:", bold=True)
TABLE(["Phương án", "Ưu / Nhược", "Quyết định"],
      [["OpenPose", "Chính xác nhưng nặng, gần như bắt buộc GPU → không hợp máy phổ thông", "Loại"],
       ["YOLO-Pose / MoveNet", "Nhanh nhưng chủ yếu 2D, ít điểm mốc, ổn định khớp kém hơn", "Loại"],
       ["MediaPipe Pose (BlazePose)", "33 điểm mốc, có tọa độ 3D, chạy thời gian thực trên CPU, miễn phí", "Chọn"]],
      widths=[2.0, 3.6, 0.8])
P("Phiên bản đầu dùng YOLOv8 làm bộ tiền lọc khung người rồi mới đưa vào MediaPipe; bản hiện "
  "tại đã bỏ YOLO, chạy MediaPipe trực tiếp trên toàn khung hình, vì bối cảnh tập tại nhà "
  "thường chỉ có một người - YOLO làm tăng độ trễ mà lợi ích không đáng kể.")
P("Tại sao tính góc 3D thay vì 2D? Bản đầu tính góc từ tọa độ 2D (x, y) trên ảnh, dễ sai khi "
  "camera đặt lệch; bản hiện tại dùng tọa độ 3D “world landmark”, tính góc bằng arccos "
  "của tích vô hướng hai vector → góc gần như độc lập với góc đặt camera (có dự phòng 2D "
  "khi thiếu dữ liệu 3D).")
P("Áp dụng: tọa độ 3 điểm (ví dụ Vai - Khuỷu - Cổ tay, hoặc Hông - Gối - Cổ chân) đưa vào "
  "công thức góc; góc này là đầu vào cho máy trạng thái đếm rep.")
H3("2.2. Xử lý tín hiệu âm thanh & NLP (phát hiện tiếng kêu đau)")
P("Là gì? Kết hợp phân loại âm thanh (nhận diện tiếng kêu/rên không lời) và nhận dạng giọng "
  "nói (bắt các từ tiếng Việt thể hiện đau).")
P("Tại sao dùng cơ chế lai YAMNet + Whisper? So sánh:", bold=True)
TABLE(["Phương án", "Vấn đề", "Quyết định"],
      [["Chỉ đo năng lượng âm (RMS)", "Biết có tiếng động to nhưng không phân biệt tiếng kêu đau với tiếng ồn", "Không đủ"],
       ["Chỉ dùng ASR (Whisper)", "Dịch tiếng la hét (Aaa, Ối) thành từ vô nghĩa → nhận nhầm nhiều", "Không đủ"],
       ["YAMNet (chính) + Whisper (phụ)", "YAMNet phân loại trực tiếp tiếng kêu; Whisper chỉ chạy khi có tiếng nói", "Chọn"]],
      widths=[2.2, 3.4, 0.8])
P("Áp dụng: mỗi đoạn âm thanh 1.5 giây tính RMS để loại khoảng lặng; nếu có tín hiệu thì "
  "YAMNet phân loại; chỉ khi YAMNet thấy có “tiếng nói” mới gọi Whisper dịch và chấm "
  "điểm từ khóa đau. Có cooldown để không đếm trùng một tiếng kêu.")
H3("2.3. Hệ chuyên gia & Máy trạng thái hữu hạn (Expert System & FSM)")
P("Là gì? FSM quản lý trạng thái hợp lệ; hệ chuyên gia dùng luật theo kinh nghiệm lâm sàng "
  "(heuristic) để ra quyết định.")
P("Tại sao dùng luật thay vì học máy? Quyết định điều chỉnh cường độ là quyết định y tế cần "
  "an toàn và GIẢI THÍCH ĐƯỢC; dữ liệu bệnh nhân nhỏ, không đủ để huấn luyện mô hình đáng tin "
  "cậy. Hệ luật cho phép kiểm soát chặt (đau nhiều → ép giảm), minh bạch lý do, dễ tinh "
  "chỉnh theo bác sĩ.")
P("Áp dụng:")
B("Đếm rep: FSM hai trạng thái Duỗi (UP) / Gập (DOWN); một rep chỉ tính khi chuyển từ Gập "
  "sang Duỗi hợp lệ → chống “ăn gian” biên độ.")
B("Brain engine: bộ phát hiện dao động dùng cửa sổ trượt 4 buổi để chặn lệnh tăng nguy hiểm.")
H3("2.4. Lập trình đồng thời & Giao tiếp tiến trình (Concurrency & IPC)")
P("Là gì? Tách các tác vụ nặng (camera/AI) khỏi giao diện bằng đa luồng; truyền hình ảnh qua "
  "một máy chủ HTTP nội bộ (MJPEG).")
P("Tại sao? Nếu xử lý AI và vẽ giao diện trên cùng một luồng, giao diện sẽ bị đơ. Tách luồng "
  "giúp UI luôn mượt kể cả trên máy cấu hình thấp.")
P("Áp dụng: một luồng chạy camera + AI và đẩy khung JPEG liên tục sang thẻ <img>; một luồng "
  "chạy micro; giao diện cập nhật số đếm + số liệu qua cơ chế gọi hàm JavaScript (evaluate_js).")
H2("3. Sơ đồ chi tiết các module (flowchart)")
P("Mã PlantUML để render bằng công cụ visual diagram.")
CODE("""
@startuml
title Luồng Thị giác & Đếm Rep (camera.py)
start
:Đọc khung hình webcam;
:Lật ngang + hiệu chỉnh sáng (gamma);
:MediaPipe Pose -> 33 landmark 3D;
if (Phát hiện người?) then (có)
  :Tính góc khớp 3D (arccos của 2 vector);
  if (Chưa "sẵn sàng"?) then (đúng)
    if (góc > up_angle liên tục >= 5 khung?) then (đủ)
      :Đánh dấu SẴN SÀNG (pha = Duỗi);
    endif
  else (đã sẵn sàng)
    if (góc < down_angle?) then (có)
      :pha = Gập; lưu góc sâu nhất + thời điểm;
    endif
    if (góc > up_angle VÀ pha trước = Gập?) then (có)
      :+1 rep; ghi ROM + tốc độ;
      :Callback updateRepCount() -> UI;
    endif
  endif
endif
:~5 lần/giây: đẩy {rep, pha, góc, đau} -> cột thông tin (updateStats);
:Vẽ khung xương + góc tại khớp; đẩy khung qua MJPEG;
stop
@enduml
""", caption="3.1. Sơ đồ luồng thị giác & đếm rep:")
CODE("""
@startuml
title Luồng Nhận diện Đau đớn (pain_detector.py)
start
:Thu micro, cắt đoạn 1.5s;
:Tính RMS (năng lượng âm);
if (RMS < ngưỡng im lặng?) then (có)
  :Bỏ qua đoạn này;
  stop
endif
:YAMNet phân loại -> cry_score, speech_score;
if (speech_score >= 0.30?) then (có tiếng nói)
  :Whisper -> văn bản tiếng Việt;
  :Chấm điểm từ khóa đau (regex) -> text_score;
endif
if (cry_score >= 0.25  HOẶC  text_score >= 2.0?) then (có)
  if (đã qua cooldown 2s?) then (có)
    :pain_count += 1;
  endif
endif
stop
@enduml
""", caption="3.2. Sơ đồ luồng nhận diện đau đớn:")
CODE("""
@startuml
title "Bộ Não" Ra Quyết Định (brain_engine.py)
start
:Nhận dữ liệu buổi tập
 (rep, tốc độ, ROM, số lần đau, tuần);
:Chấm điểm thành phần
 pain .35 / speed .25 / rom .20 / completion .20;
:Điểm tổng hợp 0-100 (+/- theo tuần);
if (đau >= 6 lần?) then (có)
  :ÉP GIẢM (an toàn) - bỏ qua bộ lọc;
  stop
endif
if (điểm >= 80?) then (có)
  :hướng = TĂNG;
elseif (điểm >= 50?) then (có)
  :hướng = GIỮ;
else (thấp)
  :hướng = GIẢM;
endif
if (đau >= 3 và đang định tăng?) then (có)
  :Hạ về GIỮ;
endif
:StabilityFilter (xác nhận buổi liên tiếp);
if (định TĂNG và mức này từng bị tụt trong 4 buổi?) then (có)
  :CHẶN tăng -> GIỮ;
endif
:RepAdjuster: tăng +1 / giảm còn 60%;
:Trả về đề xuất + lý do;
stop
@enduml
""", caption="3.3. Sơ đồ “bộ não” ra quyết định:")
H2("4. Database (sơ đồ thực thể - quan hệ)")
P("Cơ sở dữ liệu dùng SQLite, gồm 5 bảng. rehab_exercises lưu cấu hình bài tập; patient lưu "
  "hồ sơ bệnh nhân; current_config lưu cường độ hiện tại của mỗi cặp (bệnh nhân, bài tập); "
  "session_log ghi nhật ký từng buổi; exercise_adjustment ghi đề xuất điều chỉnh tương ứng.")
CODE("""
@startuml
hide circle
skinparam linetype ortho
entity rehab_exercises {
  * id : INTEGER <<PK>>
  --
  khop_tap : TEXT
  ten : TEXT
  huong_dan : TEXT
  video_url : TEXT
  lm_a / lm_b / lm_c : INTEGER
  cam_down_angle : INTEGER
  cam_up_angle : INTEGER
  ideal_sec_min / ideal_sec_max : REAL
}
entity patient {
  * patient_id : INTEGER <<PK>>
  --
  patient_name : TEXT
  date_of_birth : TEXT
  gender : TEXT
  created_at : TEXT
}
entity current_config {
  * config_id : INTEGER <<PK>>
  --
  # patient_id : INTEGER <<FK>>
  # exercise_id : INTEGER <<FK>>
  current_rep : INTEGER
  updated_at : TEXT
}
entity session_log {
  * session_log_id : INTEGER <<PK>>
  --
  # patient_id : INTEGER <<FK>>
  # exercise_id : INTEGER <<FK>>
  session_date : TEXT
  prescribed_rep : INTEGER
  actual_rep : INTEGER
  pain_count : INTEGER
}
entity exercise_adjustment {
  * adjustment_id : INTEGER <<PK>>
  --
  # session_log_id : INTEGER <<FK>>
  # patient_id : INTEGER <<FK>>
  # exercise_id : INTEGER <<FK>>
  adjustment_date : TEXT
  current_rep : INTEGER
  suggested_rep : INTEGER
  adjustment_action : TEXT
  adjustment_note : TEXT
  is_confirmed : INTEGER
  confirmed_at : TEXT
}
patient         ||--o{ current_config
rehab_exercises ||--o{ current_config
patient         ||--o{ session_log
rehab_exercises ||--o{ session_log
session_log     ||--o{ exercise_adjustment
patient         ||--o{ exercise_adjustment
rehab_exercises ||--o{ exercise_adjustment
@enduml
""", caption="Sơ đồ ERD (PlantUML):")
PB()

# =====================================================================
# III. THIET KE VA CAI DAT
# =====================================================================
H1("III. Thiết kế và cài đặt")
P("Mã nguồn tổ chức theo mô hình MVC. Phần này chỉ trình bày các class/hàm trực tiếp tạo nên "
  "logic của hệ thống; các hàm phụ trợ (mở kết nối, tải model, vẽ giao diện...) được lược bỏ. "
  "Với mỗi file: mô tả các thành phần cốt lõi, logic áp dụng, cách đánh giá, và cách kết "
  "nối - cập nhật với các file khác.", italic=True)
CODE("""
main.py                      (điểm vào -> gọi controller)
app/
  config.py                  (đường dẫn tập trung, hỗ trợ đóng gói .exe)
  controllers/api.py         (lớp Api cho UI + điều phối buổi tập)
  models/database.py         (tầng dữ liệu SQLite)
  models/brain_engine.py     (bộ não đánh giá - điều chỉnh cường độ)
  services/camera.py         (thị giác máy tính + máy chủ MJPEG)
  services/pain_detector.py  (phát hiện tiếng kêu đau)
  views/app.html             (giao diện - HTML/CSS/JS)
data/rehab.db                (database)  |  brain_states.json (trạng thái bộ não)
assets/videos/ (8 video)     |  assets/ai_models/ (model AI - tự tải)
requirements.txt  README.md  main.spec (đóng gói .exe)
""", caption="Cấu trúc thư mục (MVC):")

# ---- III.1 Database ----
H2("1. Database (models/database.py)")
P("Tầng dữ liệu dùng SQLite (thay cho SQL Server + pyodbc ở bản đầu), thiết kế theo mô hình "
  "hàm thủ tục. Các thành phần cốt lõi:")
PD("init_db()", "tạo 5 bảng nếu chưa có và chèn dữ liệu mẫu khi bảng còn rỗng. Đây là nơi đặt "
   "cấu hình 8 bài tập đã được calibrate (khớp tập, 3 điểm landmark, ngưỡng góc gập/duỗi, "
   "đường dẫn video, khoảng tốc độ chuẩn) - tức là nguồn “seed” duy nhất mà toàn bộ ứng "
   "dụng đọc ra. Hàm an toàn khi gọi nhiều lần vì dùng CREATE TABLE IF NOT EXISTS và chỉ chèn "
   "khi bảng rỗng.")
PD("save_session(...)", "ghi đồng thời hai bảng - session_log (nhật ký buổi tập) và "
   "exercise_adjustment (đề xuất điều chỉnh) - trong cùng một giao dịch. Nếu một bước lỗi thì "
   "tự động rollback toàn bộ, đảm bảo dữ liệu y tế không bị lưu nửa vời.")
PD("confirm_session(...)", "chạy khi bệnh nhân đồng ý đề xuất: đánh dấu bản ghi điều chỉnh là "
   "đã xác nhận, đồng thời cập nhật (UPSERT) số rep mới vào bảng current_config theo cặp "
   "(bệnh nhân, bài tập).")
PD("get_session_rep_info(...)", "trả về số rep hiện tại và cờ “buổi đầu” (chưa có buổi nào "
   "lưu trong session_log) - dùng để quyết định cho bệnh nhân tự chọn rep (buổi 1) hay khóa "
   "ô nhập (từ buổi 2, để AI quyết định).")
LBL("Đánh giá / ràng buộc:", "tính đúng đắn được bảo vệ ngay trong lược đồ bảng: ràng buộc "
    "UNIQUE(patient_id, exercise_id) ở current_config; các CHECK current_rep > 0, "
    "actual_rep ≥ 0, pain_count ≥ 0; cột adjustment_action chỉ nhận 'tang'/'giam'/'giu_nguyen'; "
    "các khóa ngoại liên kết bảng; bật chế độ WAL và ràng buộc khóa ngoại.")
LBL("Kết nối & cập nhật:", "bảng current_config là nguồn duy nhất về “số rep hiện tại” "
    "(brain engine không giữ giá trị này), và chỉ được cập nhật khi bệnh nhân đồng ý đề xuất. "
    "File api.py đọc cấu hình bài tập từ DB rồi đưa cho camera; bản thân camera không truy cập "
    "database trực tiếp.")

# ---- III.2 Camera ----
H2("2. Camera (services/camera.py)")
P("File đảm nhận thị giác máy tính: lấy tư thế, tính góc khớp và đếm rep. Các thành phần cốt lõi:")
PD("detect_pose()", "đưa khung hình qua mô hình MediaPipe Pose và trả về 33 điểm mốc cơ thể "
   "(cả tọa độ 2D trên ảnh lẫn tọa độ 3D thực). Mô hình chạy trên toàn khung hình, không cần "
   "bước cắt khung người bằng YOLO như bản cũ.")
PD("calc_angle_3d()", "tính góc của khớp từ tọa độ 3D, bằng arccos của tích vô hướng giữa hai "
   "vector (nối từ khớp ở giữa tới hai khớp lân cận). Vì dùng tọa độ 3D nên góc gần như không "
   "phụ thuộc vào góc đặt camera; có hàm calc_angle tính theo 2D để dự phòng.")
PD("count_rep(angle)", "là máy trạng thái (FSM) - phần logic cốt lõi của việc đếm. Trước hết "
   "buổi tập phải vào trạng thái “sẵn sàng”: góc phải vượt ngưỡng duỗi (up_angle) liên "
   "tục đủ 5 khung hình, nhằm tránh đếm nhầm lúc người mới bước vào khung. Sau khi sẵn sàng: "
   "khi góc xuống dưới ngưỡng gập (down_angle) thì chuyển sang pha Gập, đồng thời ghi lại góc "
   "sâu nhất và thời điểm bắt đầu; khi góc vượt ngưỡng duỗi trở lại mà pha trước đó là Gập thì "
   "tính một rep, lưu lại biên độ (ROM) và thời lượng của rep đó. Mỗi bài tập có ngưỡng "
   "down/up riêng, lấy từ database.")
PD("_calc_speed_score() và _calc_rom_score()", "chấm điểm chất lượng buổi tập. Điểm tốc độ so "
   "thời lượng từng rep với khoảng lý tưởng của bài (trong khoảng được 100 điểm, lệch ra ngoài "
   "thì trừ dần). Điểm ROM so góc sâu nhất từng rep với mục tiêu (down_angle) - đạt càng sâu "
   "điểm càng cao. Hai điểm này cùng số rep là dữ liệu đầu vào cho brain engine.")
PD("run()", "vòng lặp chính chạy trong một luồng riêng: đọc khung từ webcam, lật ngang và "
   "chỉnh sáng vùng tối, gọi detect_pose rồi count_rep, vẽ khung xương + số đo góc tại khớp "
   "lên khung hình rồi đẩy sang giao diện qua máy chủ MJPEG nội bộ. Đáng chú ý: thông tin buổi "
   "tập (số lần / pha / góc / số lần đau) KHÔNG còn vẽ chìm trên camera nữa mà được “đẩy” "
   "sang cột thông tin của giao diện khoảng 5 lần/giây (qua callback) - giúp khung camera "
   "thoáng và chữ hiển thị sắc nét theo chuẩn HTML.")
LBL("Kết nối & cập nhật:", "khi bắt đầu buổi, api.py tạo đối tượng đếm rep với cấu hình lấy từ "
    "DB và gắn hai callback: một bắn mỗi khi có rep mới (gọi hàm JavaScript updateRepCount để "
    "đổi số + nhấp nháy), một bắn định kỳ ~5 lần/giây (gọi updateStats để cập nhật số lần / pha "
    "/ góc / số lần đau lên cột thông tin). Giao diện nhúng thẻ <img> trỏ tới luồng video của "
    "camera, và thẻ <video> trỏ tới endpoint phục vụ clip để xem video mẫu. Khi kết thúc buổi, "
    "api.py lấy số liệu tổng kết qua get_cam_data().")

# ---- III.3 Audio ----
H2("3. Audio processing (services/pain_detector.py)")
P("File thu micro theo từng đoạn 1.5 giây (16kHz) và phát hiện đau bằng cơ chế lai hai lớp. "
  "Các thành phần cốt lõi:")
PD("SoundPainDetector.classify() (YAMNet - lớp chính)", "phân loại đoạn âm thanh, chỉ giữ các "
   "lớp “tiếng đau” (la hét, rên, khóc, gằn...) và lớp “tiếng nói”, trả về điểm "
   "tiếng kêu (cry_score) và điểm tiếng nói (speech_score).")
PD("SpeechRecognizer._process()", "điều phối xử lý một đoạn âm: tính RMS (năng lượng); nếu "
   "dưới ngưỡng im lặng thì bỏ qua cả hai mô hình; có tín hiệu thì chạy YAMNet; và CHỈ khi "
   "điểm tiếng nói đủ cao mới gọi Whisper dịch ra tiếng Việt - nhờ vậy tiếng kêu/nhiễu không "
   "bị dịch thành chữ vô nghĩa. Whisper dùng bản base nén int8 để chạy nhanh trên CPU.")
PD("score_text()", "chấm điểm đoạn văn bản bằng biểu thức chính quy: “đau” = 2.0 điểm, "
   "“đau quá” = 2.5, “ouch” = 2.0... rồi cộng lại thành tổng điểm của đoạn.")
PD("PainDetector.analyze()", "gộp hai tín hiệu: kết luận là kêu đau nếu điểm tiếng kêu ≥ 0.25 "
   "HOẶC tổng điểm từ khóa ≥ 2.0.")
PD("PainCryCounter.record()", "mỗi lần đau hợp lệ làm pain_count tăng 1, kèm cooldown 2 giây "
   "để không đếm trùng một tiếng kêu kéo dài.")
LBL("Đánh giá (ngưỡng quyết định):", "tiếng kêu tính khi cry_score ≥ 0.25; từ đau tính khi "
    "tổng điểm văn bản ≥ 2.0. Tổng số lần đau (pain_count) là tín hiệu có trọng số cao nhất "
    "(0.35) khi brain engine chấm điểm buổi tập.")
LBL("Kết nối & cập nhật:", "api.py chạy bộ thu âm trong một luồng riêng và cập nhật pain_count "
    "liên tục trong buổi; khi kết thúc, api.py đọc pain_count rồi đưa vào brain engine.")

# ---- III.4 Brain engine ----
H2("4. Brain engine (models/brain_engine.py)")
P("Hệ chuyên gia nhận dữ liệu của một buổi tập và trả về gợi ý điều chỉnh số rep. Các thành "
  "phần cốt lõi:")
PD("ScoringEngine.compute()", "chấm điểm buổi tập theo thang 0-100, bằng tổng có trọng số của "
   "bốn yếu tố (đau 0.35, tốc độ 0.25, biên độ 0.20, mức hoàn thành 0.20), rồi cộng/trừ thêm "
   "theo tuần phục hồi (giai đoạn đầu thận trọng hơn).")
PD("ScoringEngine.raw_direction()", "từ điểm số ra hướng thô: ≥ 80 → có thể TĂNG; 50-79 "
   "→ GIỮ; < 50 → GIẢM. Kèm luật an toàn: đau ≥ 6 lần ép GIẢM bất chấp điểm; đau "
   "≥ 3 lần thì không cho TĂNG.")
PD("RepAdjuster.adjust()", "tính số rep mới: tăng từng bước nhỏ (+1) nhưng giảm mạnh (xuống "
   "còn 60%) để bệnh nhân cảm nhận rõ; giới hạn trong khoảng 1-20 rep.")
PD("OscillationDetector.is_oscillating()", "quét 4 buổi gần nhất; nếu mức rep mà hệ thống định "
   "đề xuất đã từng xuất hiện rồi bị giảm xuống thì coi là đang dao động thể lực.")
PD("StabilityFilter.apply()", "lớp lọc chống dao động hai tầng: cần xác nhận buổi liên tiếp "
   "cùng hướng trước khi áp dụng thay đổi, và trước khi cho phép tăng thì gọi "
   "OscillationDetector để chặn nếu phát hiện dao động - tránh tình trạng 7→8→7→8.")
PD("BrainEngine.analyze()", "hàm tổng nối các bước: chấm điểm → phân hướng → lọc ổn "
   "định, rồi trả về một Recommendation gồm hướng điều chỉnh, số rep đề xuất và lý do.")
P("Sau khi buổi tập được lưu, hệ thống ghi số rep đã tập vào lịch sử (reps_history) và lưu "
  "xuống file brain_states.json để giữ trạng thái qua các buổi; trạng thái này tách riêng theo "
  "từng cặp (bệnh nhân, bài tập).")
TABLE(["Yếu tố", "Trọng số", "Ý nghĩa"],
      [["pain", "0.35", "Tín hiệu đau - ưu tiên cao nhất"],
       ["speed", "0.25", "Tốc độ thực hiện so với khoảng chuẩn"],
       ["rom", "0.20", "Biên độ khớp đạt được"],
       ["completion", "0.20", "Tỷ lệ hoàn thành rep mục tiêu"]],
      widths=[1.6, 1.2, 3.6], caption="Bộ trọng số chấm điểm (ScoringEngine):")
LBL("Đánh giá:", "toàn bộ bộ trọng số + ngưỡng phân hướng + luật an toàn + chống dao động ở "
    "trên chính là tiêu chí đánh giá độ phù hợp của cường độ tập cho từng bệnh nhân.")
LBL("Kết nối & cập nhật:", "api.py dựng dữ liệu đầu vào từ số liệu camera (rep, tốc độ, ROM) + "
    "số lần đau (mic) + số tuần (suy ra từ số buổi đã tập) rồi gọi analyze(); kết quả được đẩy "
    "lên giao diện. Khi bệnh nhân quyết định, api.py gọi hàm xác nhận/từ chối để cập nhật lịch "
    "sử và lưu trạng thái.")

# ---- III.5 API & UI ----
H2("5. API & UI bridge (controllers/api.py + views/app.html)")
P("Tầng điều phối nối giao diện web với toàn bộ backend qua thư viện pywebview: giao diện "
  "(JavaScript) gọi các hàm của lớp Api bằng cú pháp window.pywebview.api.<tên hàm>, mỗi hàm "
  "trả về dạng thống nhất {ok, data, error}. Các thành phần cốt lõi:")
PD("start_session()", "đọc cấu hình bài tập từ database, tạo đối tượng đếm rep, rồi áp dụng "
   "lập trình đồng thời: chạy camera và micro trên hai luồng riêng, đồng thời gắn các callback "
   "để đẩy số liệu lên giao diện theo thời gian thực.")
PD("end_session()", "dừng camera và micro, gom số liệu của buổi (số rep, tốc độ, ROM, số lần "
   "đau, số tuần), gọi brain engine phân tích và đẩy kết quả kèm đề xuất lên giao diện.")
PD("save_result()", "theo lựa chọn đồng ý/từ chối của bệnh nhân: lưu buổi tập xuống database và "
   "cập nhật số rep hiện tại (current_config) cùng trạng thái của bộ não.")
P("Xử lý số rep & ngoại lệ theo từng buổi (đặt trong end_session):", bold=True)
B("Không tập rep nào (0 rep): bỏ hẳn buổi - không lưu, không gọi brain engine, giao diện hiện "
  "thông báo “cần tập lại”.")
B("Buổi đầu của một bài mà tập chưa đủ số rep đã chọn: hạ mục tiêu bằng đúng số rep đã tập "
  "(lưu dạng “số tập được / số tập được”) để có dữ liệu nền, phần điều chỉnh còn lại do "
  "brain engine xử lý như thường.")
B("Tập đủ rep: brain engine đánh giá và điều chỉnh theo logic thông thường.")
B("Từ buổi thứ hai của một bài, ô nhập số rep bị khóa - cường độ do AI quyết định, không cho "
  "bệnh nhân tự đổi.")
P("Giao diện màn tập:", bold=True)
P("Khung camera (MJPEG) chiếm phần chính, để thoáng (chỉ còn khung xương + số đo góc tại "
  "khớp). Cột thông tin bên phải hiển thị: Số lần, Pha (Gập/Duỗi), Góc, và Số lần kêu đau; "
  "đồng hồ thời gian đặt ở thanh trên. Video hướng dẫn mặc định nằm nhỏ ở đáy cột thông tin, "
  "có thể bấm “Phóng to” để bung thành một cột riêng (và “Thu nhỏ” để thu lại); cả hai "
  "trạng thái đều chạy/dừng được. Cập nhật thời gian thực dùng evaluate_js: updateRepCount khi "
  "có rep mới, updateStats định kỳ cho số liệu sống, showResult khi kết thúc buổi.")
LBL("Kết nối & cập nhật:", "lớp Api là trung tâm điều phối: gọi database (đọc bài, lưu buổi, "
    "cập nhật rep), điều khiển camera + micro (bắt đầu/dừng/lấy số liệu), và brain engine "
    "(phân tích + xác nhận). Giao diện chỉ giao tiếp qua lớp Api, không chạm trực tiếp vào "
    "database hay các mô hình.")
PB()

# =====================================================================
# IV. SAN PHAM
# =====================================================================
H1("IV. Sản phẩm")
H2("1. Triển khai (deploy) & yêu cầu cấu hình")
P("Chạy ở chế độ phát triển: cài thư viện theo requirements.txt rồi chạy python main.py; lần "
  "đầu chạy, database tự tạo + nạp 8 bài tập, các model AI tự tải nếu thiếu.")
P("Đóng gói: ứng dụng desktop chạy cục bộ, đóng gói bằng PyInstaller (file cấu hình main.spec) "
  "thành file thực thi main.exe chạy ở chế độ cửa sổ (không hiện console). Các tài nguyên - "
  "giao diện app.html, model pose, model YAMNet và thư mục video - đều được nhúng kèm; file ghi "
  "được (rehab.db, brain_states.json) đặt cạnh file .exe. Nhờ nhúng sẵn cả hai model, bản .exe "
  "chạy được offline hoàn toàn (không cần tải gì thêm). Không có máy chủ hay cloud - dữ liệu "
  "bệnh nhân nằm hoàn toàn trên máy người dùng.")
TABLE(["Thành phần", "Tối thiểu", "Khuyến nghị"],
      [["Hệ điều hành", "Windows 10/11 64-bit", "Windows 11"],
       ["CPU", "2 nhân (chạy MediaPipe + YAMNet + Whisper base trên CPU)", "4 nhân trở lên"],
       ["RAM", "4 GB", "8 GB"],
       ["Thiết bị", "Webcam 720p + micro", "Webcam 1080p"],
       ["Ổ đĩa / mạng", "~200 MB; bản .exe chạy offline, bản dev cần mạng tải model lần đầu", "-"],
       ["GPU", "Không bắt buộc", "Không cần"]],
      widths=[1.7, 3.4, 1.4])
H2("2. Thông báo / phản hồi cho người dùng")
B("Thời gian thực: số lần (rep) nhảy ngay sau mỗi lần tập (kèm nhấp nháy), kèm Pha và Góc cập "
  "nhật liên tục ở cột thông tin; đồng hồ buổi tập ở thanh trên.")
B("Cảnh báo an toàn: khi phát hiện tiếng kêu đau, hệ thống ghi nhận và đếm (hiển thị Số lần "
  "kêu đau); số lần đau tác động trực tiếp tới đề xuất cường độ (đau nhiều → ép giảm).")
B("Kết thúc buổi: bảng kết quả (rep đạt / mục tiêu, số lần đau, thời lượng, điểm) + đề xuất "
  "tăng/giảm/giữ kèm lý do để bệnh nhân quyết định; nếu không tập rep nào thì hiện “cần tập lại”.")
B("Hướng phát triển: thêm nhắc nghỉ giữa hiệp, cảnh báo âm thanh, thông báo cho bác sĩ theo "
  "dõi từ xa.")
H2("3. Dữ liệu & cách tạo bài tập")
P("Về mô hình AI: dự án KHÔNG tự huấn luyện mô hình mới mà dùng ba mô hình tiền huấn luyện "
  "(MediaPipe Pose, YAMNet, Whisper). Do đó “dữ liệu” của dự án không phải tập huấn "
  "luyện, mà gồm ba loại:")
N("Tập video hướng dẫn tự quay: 8 clip (1 clip / bài), vừa làm video mẫu, vừa là nguồn để "
  "calibrate ngưỡng góc.")
N("Dữ liệu seed trong database: cấu hình 8 bài tập (khớp, landmark, ngưỡng góc, tốc độ chuẩn) "
  "+ bệnh nhân mẫu.")
N("Dữ liệu vận hành sinh ra khi dùng: session_log và exercise_adjustment - lịch sử buổi tập, "
  "được brain engine dùng để chống dao động và điều chỉnh cường độ.")
P("Cách tạo một bài tập (calibrate):", bold=True)
N("Quay video thực hiện đúng động tác (chuyển .MOV → .mp4).")
N("Chạy MediaPipe trên video để đo biên độ góc 3D (min → max) của khớp trong cả chu kỳ.")
N("Đặt ngưỡng theo “band giữa” (lùi ~30% mỗi đầu của biên độ) để rep tính được với "
  "bệnh nhân ROM hạn chế, không đòi gập/duỗi kịch.")
N("Ghi ngưỡng vào database (cột cam_down_angle / cam_up_angle).")
P("Tiêu chí phù hợp của bài tập:", bold=True)
B("Phù hợp với khớp và giai đoạn phục hồi (gối ACL / khuỷu Tennis Elbow; đầu / giữa / cuối).")
B("Ngưỡng góc đạt được với biên độ thực tế của bệnh nhân (đã calibrate theo band giữa).")
B("Cường độ (số rep) được brain engine đánh giá lại mỗi buổi theo điểm đau/tốc độ/ROM/hoàn "
  "thành, có ràng buộc an toàn.")
P("8 bài tập hiện có và ngưỡng góc (calibrate từ video):", bold=True)
TABLE(["#", "Tên bài", "Khớp", "down / up (°)"],
      [["1", "trượt gối", "gối", "94 / 133"], ["2", "nâng chân thẳng", "gối", "135 / 166"],
       ["3", "ngồi dựa tường", "gối", "109 / 145"], ["4", "gập gối đứng", "gối", "100 / 140"],
       ["5", "gập/duỗi khuỷu tay", "khuỷu", "66 / 129"], ["6", "duỗi tay trên đầu", "khuỷu", "84 / 130"],
       ["7", "gập cánh tay đứng", "khuỷu", "84 / 127"], ["8", "duỗi khuỷu nhờ trọng lực", "khuỷu", "56 / 122"]],
      widths=[0.5, 2.8, 1.0, 1.7])
H2("4. Phương pháp đánh giá")
P("Hệ thống được đánh giá theo từng module bằng thử nghiệm thực tế (chưa phải thử nghiệm lâm sàng):")
B("Đếm rep: so số rep app đếm với đếm thủ công trên cùng động tác; đo tỷ lệ đếm thiếu / nhầm "
  "và tinh chỉnh ngưỡng. Ví dụ: bài “nâng chân thẳng” ban đầu band hẹp (19°) gây đếm "
  "nhầm khi rung → đã mở rộng band lên 31°.")
B("Phát hiện đau: thử với kêu đau / nói từ “đau” / im lặng / nói chuyện thường, đo tỷ lệ "
  "phát hiện đúng và báo nhầm; tinh chỉnh ngưỡng YAMNet và cooldown.")
B("Đo góc: kiểm tra độ ổn định của góc 3D khi đổi góc đặt camera (so với 2D).")
B("Hiệu năng: đo FPS và độ trễ rep → giao diện trên máy mục tiêu (CPU); xác nhận UI không "
  "đơ nhờ tách luồng.")
B("Brain engine: kiểm thử logic chấm điểm và chống dao động bằng kịch bản giả lập (lịch sử "
  "rep, số lần đau khác nhau).")
P("Hạn chế của đánh giá: quy mô nhỏ, chưa thử nghiệm trên bệnh nhân thật và chưa đối chiếu với "
  "đánh giá của kỹ thuật viên vật lý trị liệu.", italic=True)
H2("5. Demo sản phẩm")
P("[Chèn ảnh chụp màn hình: đăng nhập → chọn bài (lưới 8 bài) → màn tập (camera + cột "
  "thông tin + video mẫu) → màn kết quả đề xuất.]", italic=True)
PB()

# =====================================================================
# V. KET LUAN
# =====================================================================
H1("V. Kết luận")
P("Nhóm đã xây dựng được một hệ thống AI hỗ trợ phục hồi chức năng chạy cục bộ, kết hợp ba "
  "mảng: thị giác máy tính (đo góc khớp 3D, đếm rep), xử lý âm thanh (phát hiện tiếng kêu đau "
  "bằng YAMNet + Whisper), và hệ chuyên gia (đánh giá - điều chỉnh cường độ an toàn). Hệ thống "
  "có cơ sở dữ liệu SQLite, giao diện trực quan với video hướng dẫn, và đóng gói được thành "
  "ứng dụng desktop chạy offline.")
P("Hạn chế:", bold=True)
B("Chỉ theo dõi tốt khi có một người trong khung hình; nhạy với ánh sáng và góc đặt camera.")
B("Dùng model tiền huấn luyện, chưa tinh chỉnh cho dữ liệu phục hồi chức năng chuyên biệt.")
B("Chưa thử nghiệm lâm sàng; ngưỡng góc calibrate từ video mẫu nên phụ thuộc chất lượng clip.")
P("Hướng phát triển:", bold=True)
B("Mở rộng thêm khớp và bài tập; thêm bài đếm thời gian giữ tư thế.")
B("Xây dựng tập dữ liệu có nhãn để tinh chỉnh/đánh giá mô hình bài bản hơn.")
B("Tùy chọn đồng bộ đám mây và bảng theo dõi cho bác sĩ; phiên bản di động.")
B("Thử nghiệm thực tế với bệnh nhân dưới giám sát của chuyên gia.")

doc.save(OUT)
print("SAVED:", OUT)
